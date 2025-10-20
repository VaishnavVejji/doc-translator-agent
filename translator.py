from io import BytesIO
from docx import Document
from deep_translator import GoogleTranslator
import fitz  # PyMuPDF

def translate_text(text, target_language):
    if not text.strip():
        return ""
    try:
        return GoogleTranslator(source="auto", target=target_language).translate(text)
    except Exception:
        return text

# ---------- DOCX ----------
def translate_docx(input_stream, target_language):
    input_stream.seek(0)
    doc = Document(input_stream)
    translated_doc = Document()
    translated_text_full = []

    for para in doc.paragraphs:
        translated_line = translate_text(para.text, target_language)
        translated_doc.add_paragraph(translated_line)
        translated_text_full.append(translated_line)

    output_stream = BytesIO()
    translated_doc.save(output_stream)
    output_stream.seek(0)
    return output_stream, "\n".join(translated_text_full)

# ---------- PDF ----------
def extract_and_translate(input_stream, target_language):
    input_stream.seek(0)
    pdf_data = input_stream.read()
    if not pdf_data:
        raise ValueError("Empty PDF file received.")

    doc = fitz.open(stream=pdf_data, filetype="pdf")
    translated_lines = []

    # Translate line by line
    for page in doc:
        lines = page.get_text("text").splitlines()
        for line in lines:
            if line.strip():
                translated_lines.append(translate_text(line, target_language))

    # Combine for preview
    translated_text_full = "\n".join(translated_lines)

    # Generate translated PDF
    output_pdf = fitz.open()
    for page in doc:
        new_page = output_pdf.new_page(width=page.rect.width, height=page.rect.height)
        blocks = page.get_text("blocks") or []
        for block in blocks:
            if len(block) >= 5:
                x0, y0, x1, y1, text = block[:5]
                if text.strip():
                    translated_block = translate_text(text, target_language)
                    new_page.insert_text((x0, y0), translated_block, fontsize=11)

    pdf_out = BytesIO(output_pdf.write())
    pdf_out.seek(0)
    return pdf_out, translated_text_full
